from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_border(cell, **kwargs):
    """
    Set cell border
    Usage:
    set_cell_border(
        cell,
        top={"sz": 12, "val": "single", "color": "#000000", "space": "0"},
        bottom={"sz": 12, "color": "#000000", "val": "single"},
        start={"sz": 24, "val": "dashed", "color": "#000000"},
        end={"sz": 12, "val": "dashed", "color": "#000000"}
    )
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # check for tag existance, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # list the sides
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            
            # check for tag existance, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)

            # look for the attribute and set it
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def parse_pos_text(text):
    """Parse the input text and extract POS tagging data."""
    lines = text.split('\n')
    parsed_data = []
    current_sentence = ''
    
    for line in lines:
        # Skip empty lines and separator lines
        if not line.strip() or line.startswith('-----') or line.startswith('Token') or line.startswith('Sentence'):
            continue
            
        # Capture full sentence
        if line.startswith('Full sentence:'):
            current_sentence = line.replace('Full sentence:', '').strip()
            continue
            
        # Parse data lines
        parts = line.strip().split()
        if len(parts) >= 4:
            parsed_data.append({
                'token': parts[0],
                'true_label': parts[1],
                'predicted_label': parts[2],
                'match': parts[3]
            })
    
    return parsed_data, current_sentence

def create_word_document(parsed_data, sentence, output_filename):
    """Create a Word document with the POS tagging results in a table."""
    doc = Document()
    
    # Add title
    title = doc.add_paragraph('Part of Speech (POS) Tagging Results')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(14)
    title.runs[0].font.bold = True
    
    # Add table
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    
    # Set header row
    header_cells = table.rows[0].cells
    headers = ['Word Id', 'Word', 'True Label', 'Predicted Label', 'Match?']
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].font.bold = True
        
        # Add borders to header cells
        border_settings = {
            "top": {"sz": 12, "val": "single", "color": "#000000"},
            "bottom": {"sz": 12, "val": "single", "color": "#000000"},
            "start": {"sz": 12, "val": "single", "color": "#000000"},
            "end": {"sz": 12, "val": "single", "color": "#000000"}
        }
        set_cell_border(header_cells[i], **border_settings)
    
    # Add data rows
    for item in parsed_data:
        row_cells = table.add_row().cells
        row_cells[1].text = item['token']
        row_cells[2].text = item['true_label']
        row_cells[3].text = item['predicted_label']
        row_cells[4].text = item['match']
        
        # Add borders to data cells
        for cell in row_cells:
            set_cell_border(cell, **border_settings)
    
    # Set column widths
    for column in table.columns:
        for cell in column.cells:
            cell.width = Inches(1.5)
    
    # Add full sentence
    doc.add_paragraph()  # Add spacing
    sentence_para = doc.add_paragraph('Full sentence: ')
    sentence_para.add_run(sentence).italic = True
    
    # Save the document
    doc.save(output_filename)

def convert_pos_to_word(input_text, output_filename):
    """Main function to convert POS text to Word document."""
    # Parse the input text
    parsed_data, sentence = parse_pos_text(input_text)
    
    # Create Word document
    create_word_document(parsed_data, sentence, output_filename)
    
    print(f"Word document created successfully: {output_filename}")

# Example usage
if __name__ == "__main__":
    # Sample input text
    sample_text = """Sentence 6:
------------------------------------------------------------
Token           True Label      Predicted Label Match?
------------------------------------------------------------
Ang             DET             DET             ✓
hukbo           NN              NN              ✓
sang            CONJ            CONJ            ✓
Israel          NN              NN              ✓
ginatawag       VB              VB              ✓
ang             DET             DET             ✓
bag-o           ADJ             ADJ             ✓
grupo           NN              NUM             ✗
sang            CONJ            CONJ            ✓
mga             DET             NN              ✗
barko           NN              CONJ            ✗
nga             CONJ            DET             ✗
"               PUNCT           NN              ✗
iloy            NN              CONJ            ✗
sang            CONJ            PUNCT           ✗
tanan           ADJ             NN              ✗
nga             CONJ            CONJ            ✓
mga             DET             PRNN            ✗
plotilya        NN              CONJ            ✗
.               PUNCT           DET             ✗
"               PUNCT           NN              ✗
"""

    # Convert to Word document
    convert_pos_to_word(sample_text, 'pos_tagging_results.docx')